from __future__ import annotations

import json
from pathlib import Path

from PIL import Image
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT_DIR = Path(__file__).resolve().parents[1]
EVIDENCE_DIR = ROOT_DIR / "evidencias" / "actividad_6"
ASSET_DIR = EVIDENCE_DIR / "report_assets"
OUTPUT_DIR = ROOT_DIR / "entregables" / "actividad_6"
OUTPUT_FILE = OUTPUT_DIR / "Informe_Actividad_6_Selenium.docx"
RESULTS_FILE = EVIDENCE_DIR / "resultados_pruebas.json"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "0B2545"
RED = "C0392B"
GREEN = "1F6F43"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "667085"
WHITE = "FFFFFF"
BLACK = "111111"
TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def rgb(hex_color: str) -> RGBColor:
    return RGBColor.from_string(hex_color)


def set_run_font(
    run,
    *,
    name: str = "Calibri",
    size: float = 11,
    color: str = BLACK,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="D0D5DD", size="6") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa: list[int]) -> None:
    if sum(widths_dxa) != TABLE_WIDTH_DXA:
        raise ValueError(f"Las columnas deben sumar {TABLE_WIDTH_DXA} DXA")

    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(TABLE_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths_dxa[index]
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)

    set_table_borders(table)


def style_table_text(table, header=True, center_columns: set[int] | None = None) -> None:
    center_columns = center_columns or set()
    for row_index, row in enumerate(table.rows):
        for column_index, cell in enumerate(row.cells):
            is_header = header and row_index == 0
            if is_header:
                shade_cell(cell, NAVY)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.05
                paragraph.alignment = (
                    WD_ALIGN_PARAGRAPH.CENTER
                    if column_index in center_columns
                    else WD_ALIGN_PARAGRAPH.LEFT
                )
                for run in paragraph.runs:
                    set_run_font(
                        run,
                        size=9 if is_header else 9.2,
                        color=WHITE if is_header else BLACK,
                        bold=is_header,
                    )
    if header:
        set_repeat_table_header(table.rows[0])


def add_table(doc, headers: list[str], rows: list[list[str]], widths_dxa: list[int], center=None):
    table = doc.add_table(rows=1, cols=len(headers))
    for index, value in enumerate(headers):
        table.rows[0].cells[index].text = value
    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = str(value)
    set_table_geometry(table, widths_dxa)
    style_table_text(table, center_columns=set(center or []))
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)
    return table


def set_paragraph_shading(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def add_labeled_paragraph(doc, label: str, text: str, *, after=6):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(after)
    lead = paragraph.add_run(f"{label}: ")
    set_run_font(lead, bold=True, color=NAVY)
    body = paragraph.add_run(text)
    set_run_font(body)
    return paragraph


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Página ")
    set_run_font(run, size=9, color=MID_GRAY)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, text, end])


def add_hyperlink(paragraph, text: str, url: str) -> None:
    part = paragraph.part
    relationship_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.extend([color, underline])
    text_element = OxmlElement("w:t")
    text_element.text = text
    run.extend([r_pr, text_element])
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(BLACK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    specs = {
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in specs.items():
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    caption = doc.styles["Caption"]
    caption.font.name = "Calibri"
    caption._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    caption._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = rgb(MID_GRAY)
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(8)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER


def configure_page(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    header = section.header
    header_paragraph = header.paragraphs[0]
    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header_paragraph.paragraph_format.space_after = Pt(0)
    run = header_paragraph.add_run("SUTRAN VIAL  |  ACTIVIDAD 6")
    set_run_font(run, size=9, color=MID_GRAY, bold=True)

    add_page_number(section.footer.paragraphs[0])


def add_cover(doc: Document, results: dict) -> None:
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(55)

    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_after = Pt(10)
    run = kicker.add_run("CALIDAD DE SOFTWARE")
    set_run_font(run, size=10, color=RED, bold=True)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(7)
    run = title.add_run("Actividad 6")
    set_run_font(run, size=30, color=NAVY, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(24)
    run = subtitle.add_run("Evaluación automatizada de calidad con Selenium")
    set_run_font(run, size=16, color=DARK_BLUE, bold=True)

    project = doc.add_paragraph()
    project.paragraph_format.space_after = Pt(26)
    run = project.add_run(
        "Plataforma Web para el Monitoreo y Visualización de Accidentes de Tránsito "
        "en Carreteras del Perú basada en datos abiertos de SUTRAN, periodo 2020-2021"
    )
    set_run_font(run, size=12, color=MID_GRAY)

    add_labeled_paragraph(doc, "Herramienta", "Selenium WebDriver 4.45.0")
    add_labeled_paragraph(doc, "Fecha de ejecución", "30 de junio de 2026")
    add_labeled_paragraph(doc, "Entorno", "Python 3.14.2, Google Chrome 149, Windows")

    summary = doc.add_paragraph()
    summary.paragraph_format.space_before = Pt(22)
    summary.paragraph_format.space_after = Pt(10)
    summary.paragraph_format.left_indent = Inches(0.12)
    summary.paragraph_format.right_indent = Inches(0.12)
    summary.paragraph_format.space_before = Pt(10)
    summary.paragraph_format.space_after = Pt(10)
    set_paragraph_shading(summary, "EAF7EF")
    lead = summary.add_run("RESULTADO FINAL  ")
    set_run_font(lead, size=11, color=GREEN, bold=True)
    value = summary.add_run(
        f"{results['aprobadas']} de {results['total_pruebas']} pruebas aprobadas  |  "
        f"{results['tasa_exito_porcentaje']:.0f}% de éxito  |  "
        f"{results['duracion_total_segundos']:.3f} s"
    )
    set_run_font(value, size=11, color=GREEN, bold=True)

    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(18)
    run = note.add_run("Informe técnico de ejecución, evidencias, análisis y diagnóstico.")
    set_run_font(run, size=10, color=MID_GRAY, italic=True)

    doc.add_page_break()


def add_intro_and_strategy(doc: Document, results: dict) -> None:
    doc.add_heading("1. Objetivo y alcance", level=1)
    doc.add_paragraph(
        "La Actividad 6 verifica la calidad funcional de la plataforma SUTRAN VIAL mediante "
        "pruebas automatizadas sobre la lógica de datos, la API REST y la interfaz web. El "
        "alcance incluye la carga del dataset oficial, filtros combinados, indicadores, mapa "
        "departamental, tabla de registros y exportación del reporte Excel."
    )
    doc.add_paragraph(
        "La evaluación toma como referencia el modelo de calidad de producto ISO/IEC "
        "25010:2023 para organizar el diagnóstico y emplea WebDriver como mecanismo de "
        "automatización del navegador."
    )

    doc.add_heading("2. Herramienta instalada", level=1)
    add_labeled_paragraph(doc, "Herramienta seleccionada", "Selenium WebDriver 4.45.0")
    add_labeled_paragraph(
        doc,
        "Instalación",
        "Se instaló en el entorno virtual del proyecto mediante el comando "
        "python -m pip install -U selenium.",
    )
    add_labeled_paragraph(
        doc,
        "Controlador",
        "Selenium Manager administra el controlador compatible con Google Chrome, evitando "
        "una configuración manual del ejecutable.",
    )

    add_table(
        doc,
        ["Componente", "Versión / configuración"],
        [
            ["Sistema operativo", "Microsoft Windows NT 10.0.26200.0"],
            ["Python", "3.14.2"],
            ["Selenium", "4.45.0"],
            ["Google Chrome", "149.0.7827.197, modo headless"],
            ["Backend", "Servidor Python, http://127.0.0.1:8000"],
            ["Frontend", "Servidor HTTP, http://127.0.0.1:5500"],
            ["Dataset", "SUTRAN 2020-2021, 8,155 registros"],
        ],
        [2520, 6840],
    )

    doc.add_heading("3. Estrategia de pruebas", level=1)
    doc.add_paragraph(
        "Se aplicó una estrategia por niveles para detectar fallos cerca de su origen y, al "
        "mismo tiempo, comprobar los recorridos visibles para un usuario final."
    )
    add_table(
        doc,
        ["Nivel", "Cantidad", "Propósito"],
        [
            ["Unitarias", "6", "Validar carga, normalización, filtros y KPI del servicio de datos."],
            ["Integración API", "5", "Comprobar contrato HTTP, opciones, filtros y validaciones."],
            ["Funcionales Selenium", "6", "Validar navegación, interfaz, mapa, tabla y exportación."],
        ],
        [2160, 1080, 6120],
        center=[1],
    )
    doc.add_paragraph(
        f"La batería completa contiene {results['total_pruebas']} casos y se ejecuta con un "
        "único comando reproducible. Las pruebas Selenium generan capturas PNG y verifican "
        "la descarga de un archivo Excel filtrado."
    )


def add_test_matrix(doc: Document) -> None:
    doc.add_heading("4. Casos de prueba implementados", level=1)
    cases = [
        ["UT-01", "Unitaria", "Cargar exactamente 8,155 registros del CSV", "Aprobada"],
        ["UT-02", "Unitaria", "Reconocer 25 departamentos e incluir Callao con 6 casos", "Aprobada"],
        ["UT-03", "Unitaria", "Filtrar exclusivamente los registros de Callao", "Aprobada"],
        ["UT-04", "Unitaria", "Combinar año 2021 y modalidad Choque", "Aprobada"],
        ["UT-05", "Unitaria", "Mantener consistencia entre KPI y filas filtradas", "Aprobada"],
        ["UT-06", "Unitaria", "Responder vacío ante un departamento inexistente", "Aprobada"],
        ["API-01", "Integración", "Responder estado saludable en /api/health", "Aprobada"],
        ["API-02", "Integración", "Exponer 25 departamentos en opciones", "Aprobada"],
        ["API-03", "Integración", "Devolver resumen de Callao con 6 accidentes", "Aprobada"],
        ["API-04", "Integración", "Devolver 6 registros coherentes de Callao", "Aprobada"],
        ["API-05", "Integración", "Rechazar el año inválido 2019 con HTTP 422", "Aprobada"],
        ["E2E-01", "Funcional", "Cargar la portada y los indicadores reales", "Aprobada"],
        ["E2E-02", "Funcional", "Navegar desde la portada hasta el dashboard", "Aprobada"],
        ["E2E-03", "Funcional", "Filtrar Callao y colorear una sola geometría", "Aprobada"],
        ["E2E-04", "Funcional", "Mostrar 6 filas de Callao en la tabla", "Aprobada"],
        ["E2E-05", "Funcional", "Combinar filtros de año y modalidad", "Aprobada"],
        ["E2E-06", "Funcional", "Previsualizar y descargar el Excel filtrado", "Aprobada"],
    ]
    add_table(
        doc,
        ["ID", "Tipo", "Escenario verificado", "Resultado"],
        cases,
        [900, 1440, 5580, 1440],
        center=[0, 1, 3],
    )


def add_results_and_diagnosis(doc: Document, results: dict) -> None:
    doc.add_heading("5. Resultados obtenidos", level=1)
    add_table(
        doc,
        ["Indicador", "Resultado"],
        [
            ["Pruebas ejecutadas", str(results["total_pruebas"])],
            ["Pruebas aprobadas", str(results["aprobadas"])],
            ["Pruebas fallidas", str(results["fallidas"])],
            ["Errores de ejecución", str(results["errores"])],
            ["Pruebas omitidas", str(results["omitidas"])],
            ["Tasa de éxito", f"{results['tasa_exito_porcentaje']:.0f}%"],
            ["Duración total", f"{results['duracion_total_segundos']:.3f} segundos"],
            ["Capturas generadas", str(len(results["capturas"]))],
        ],
        [4680, 4680],
        center=[1],
    )
    doc.add_paragraph(
        "La ejecución final fue satisfactoria: no se detectaron fallos funcionales en los "
        "escenarios cubiertos. La API mantuvo coherencia con el dataset y la interfaz reflejó "
        "correctamente los filtros, incluido Callao."
    )

    doc.add_heading("6. Evaluación de calidad", level=1)
    add_table(
        doc,
        ["Característica", "Evidencia observada", "Evaluación"],
        [
            ["Adecuación funcional", "Filtros, KPI, mapa, tabla y exportación aprobados.", "Satisfactoria"],
            ["Fiabilidad", "17 casos repetibles sin fallos en la ejecución final.", "Satisfactoria en entorno local"],
            ["Compatibilidad", "Pruebas ejecutadas en Google Chrome 149.", "Parcial: falta Edge/Firefox"],
            ["Usabilidad", "Navegación y respuesta visual verificadas por recorridos E2E.", "Satisfactoria en flujos cubiertos"],
            ["Mantenibilidad", "Pruebas separadas por unidad, API y E2E.", "Base adecuada para regresión"],
            ["Eficiencia", "La suite completa tardó 13.758 segundos.", "Sin prueba formal de carga"],
            ["Seguridad", "Validación de año inválido comprobada.", "Evaluación limitada"],
        ],
        [1980, 4680, 2700],
    )

    doc.add_heading("7. Análisis y diagnóstico", level=1)
    add_labeled_paragraph(
        doc,
        "Diagnóstico principal",
        "La plataforma es funcionalmente estable dentro del alcance evaluado. Los resultados "
        "del frontend, la API y el servicio de datos son consistentes entre sí.",
    )
    add_labeled_paragraph(
        doc,
        "Incidencia durante la preparación",
        "La primera ejecución obtuvo 16 de 17 pruebas aprobadas porque el script interpretó "
        "el valor localizado 4,145 con una regla numérica incorrecta. Se clasificó como defecto "
        "del código de prueba, no del producto; se corrigió la sincronización y la suite final "
        "alcanzó 100% de éxito.",
    )
    add_labeled_paragraph(
        doc,
        "Riesgo externo",
        "Los mapas y bibliotecas visuales dependen de recursos remotos. Una interrupción de "
        "Internet podría afectar pruebas y visualización aun cuando el backend local funcione.",
    )
    add_labeled_paragraph(
        doc,
        "Cobertura pendiente",
        "No se realizaron pruebas de carga, penetración, accesibilidad, recuperación ante fallos "
        "ni ejecución cruzada completa en otros navegadores.",
    )

    doc.add_heading("8. Recomendaciones", level=1)
    add_table(
        doc,
        ["Prioridad", "Acción recomendada", "Beneficio esperado"],
        [
            ["Alta", "Ejecutar la suite en cada cambio mediante integración continua.", "Evitar regresiones."],
            ["Alta", "Incorporar Edge y Firefox a las pruebas E2E.", "Ampliar compatibilidad."],
            ["Media", "Alojar localmente una copia versionada del GeoJSON.", "Reducir dependencia externa."],
            ["Media", "Agregar pruebas de carga a los endpoints de consulta.", "Medir eficiencia y capacidad."],
            ["Media", "Añadir revisión de accesibilidad WCAG.", "Mejorar inclusión y usabilidad."],
            ["Baja", "Ampliar validaciones negativas de parámetros y exportación.", "Fortalecer robustez."],
        ],
        [1260, 4860, 3240],
        center=[0],
    )


def crop_for_report(source: Path, output: Path, bottom: int | None = None) -> Path:
    with Image.open(source) as image:
        if bottom is None:
            cropped = image.copy()
        else:
            cropped = image.crop((0, 0, image.width, min(bottom, image.height)))
        cropped.save(output)
    return output


def add_figure(doc: Document, path: Path, caption: str, width=6.25) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run()
    run.add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph(caption, style="Caption")
    cap.paragraph_format.keep_with_next = False


def add_evidence(doc: Document) -> None:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    assets = [
        ("01_pagina_principal.png", "01_portada.png", None, "Figura 1. Página principal con 25 departamentos y 8,155 accidentes."),
        ("02_dashboard_resumen.png", "02_dashboard.png", None, "Figura 2. Dashboard analítico con KPI, gráficos y mapa de calor temporal."),
        ("03_filtro_callao_mapa.png", "03_callao_mapa.png", 650, "Figura 3. Filtro Callao: una sola geometría resaltada y 6 registros."),
        ("04_tabla_callao.png", "04_callao_tabla.png", 400, "Figura 4. Tabla de datos con los 6 accidentes correspondientes a Callao."),
        ("05_reporte_callao.png", "05_callao_reporte.png", 450, "Figura 5. Previsualización del reporte Excel filtrado para Callao."),
    ]

    doc.add_heading("9. Evidencias de ejecución", level=1)
    doc.add_paragraph(
        "Las siguientes capturas fueron generadas automáticamente por Selenium durante la "
        "ejecución satisfactoria de la batería de pruebas."
    )
    for index, (source_name, output_name, bottom, caption) in enumerate(assets):
        if index:
            doc.add_page_break()
        asset = crop_for_report(
            EVIDENCE_DIR / source_name, ASSET_DIR / output_name, bottom=bottom
        )
        add_figure(doc, asset, caption)


def add_reproducibility_and_references(doc: Document) -> None:
    doc.add_page_break()
    doc.add_heading("10. Reproducción de la práctica", level=1)
    doc.add_paragraph(
        "Con el backend y el frontend activos, la batería completa se ejecuta desde la raíz "
        "del proyecto con el siguiente comando:"
    )
    code = doc.add_paragraph()
    code.paragraph_format.left_indent = Inches(0.25)
    code.paragraph_format.right_indent = Inches(0.25)
    code.paragraph_format.space_before = Pt(4)
    code.paragraph_format.space_after = Pt(8)
    set_paragraph_shading(code, LIGHT_GRAY)
    run = code.add_run("python tests\\run_activity6.py")
    set_run_font(run, name="Consolas", size=10, color=NAVY, bold=True)

    add_table(
        doc,
        ["Artefacto", "Ubicación"],
        [
            ["Pruebas unitarias", "tests/test_data_service.py"],
            ["Pruebas de API", "tests/test_api.py"],
            ["Pruebas Selenium", "tests/test_selenium.py"],
            ["Ejecutor y resultados", "tests/run_activity6.py"],
            ["Evidencias", "evidencias/actividad_6/"],
            ["Dependencias", "requirements-test.txt"],
        ],
        [2880, 6480],
    )

    doc.add_heading("11. Conclusiones", level=1)
    doc.add_paragraph(
        "La instalación de Selenium y la implementación de pruebas en tres niveles permitieron "
        "evaluar el comportamiento del proyecto con evidencia objetiva y repetible. La ejecución "
        "final alcanzó 17 de 17 pruebas aprobadas, sin fallos funcionales en el alcance definido."
    )
    doc.add_paragraph(
        "El diagnóstico es favorable para una presentación académica y una demostración local. "
        "La siguiente etapa de madurez debe ampliar la cobertura a compatibilidad cruzada, "
        "rendimiento, accesibilidad, seguridad e integración continua."
    )

    doc.add_heading("12. Referencias", level=1)
    references = [
        ("Selenium WebDriver - Getting Started", "https://www.selenium.dev/documentation/webdriver/getting_started/"),
        ("Selenium Python API 4.45.0", "https://www.selenium.dev/selenium/docs/api/py/"),
        ("ISO/IEC 25010:2023 - Product quality model", "https://www.iso.org/standard/78176.html"),
        ("W3C WebDriver", "https://www.w3.org/TR/webdriver/"),
    ]
    for label, url in references:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(5)
        add_hyperlink(paragraph, label, url)


def build() -> Path:
    if not RESULTS_FILE.exists():
        raise FileNotFoundError(
            "Primero ejecute tests/run_activity6.py para generar resultados reales."
        )
    results = json.loads(RESULTS_FILE.read_text(encoding="utf-8"))
    if results.get("tasa_exito_porcentaje") != 100.0:
        raise RuntimeError("El informe final requiere una ejecución satisfactoria.")

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_styles(doc)
    configure_page(doc)
    add_cover(doc, results)
    add_intro_and_strategy(doc, results)
    add_test_matrix(doc)
    add_results_and_diagnosis(doc, results)
    add_evidence(doc)
    add_reproducibility_and_references(doc)

    properties = doc.core_properties
    properties.title = "Actividad 6 - Evaluación de calidad con Selenium"
    properties.subject = "Pruebas de software, resultados, diagnóstico y evidencias"
    properties.author = "Equipo del proyecto SUTRAN VIAL"
    properties.keywords = "Selenium, pruebas de software, calidad, SUTRAN"

    doc.save(OUTPUT_FILE)
    return OUTPUT_FILE


if __name__ == "__main__":
    print(build())
